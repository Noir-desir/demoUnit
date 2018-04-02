#! usr/bin/python
# _*_ coding=utf-8 _*_
'''
@time: 2018/3/27 17:05
@author: jiangzeyu5
'''

from docx import Document
import subprocess
import re
import os

class Doc_handle():
    def __init__(self):
        self._doc = None
        self.conn_doc()

    def conn_doc(self):
        for filename in os.listdir(os.getcwd()):
            # if filename.endswith('.doc'):
            #     subprocess.call(['soffice', '--headless', '--convert-to', 'docx', filename])
            #     self._doc = Document(filename[:-4]+'.docx')
            if filename.endswith('.docx'):
                self._doc = Document(filename)
            else:
                None

    def param_text(self, regex_rules):
        results = [ ]
        pattern = re.compile(regex_rules, re.I)
        for para in self._doc.paragraphs:
            m = pattern.findall(para.text)
            if m:
                a ="".join(m)
                results.append(a)
        print(results)
        return results  # 返回字符串列表

    def param_table(self, regex_rules):
        results = []
        pattern = re.compile(regex_rules, re.I)
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    m = pattern.findall(cell.text)
                    if m:
                        a = "".join(m)
                        results.append(a)
        print(results)
        return results  # 返回字符串列表

    def save_text(self, textname, result):
        filepath = textname + '.txt'
        with open(filepath, 'a', encoding='utf-8') as f:
            f.write(result+'\n')

if __name__ == '__main__':
    jiexi = Doc_handle().param_text(r'PR-F.*')
    # jiexi = Doc_handle().param_table(r'PR.*')

    for result in jiexi:
        Doc_handle().save_text('haha',result)
